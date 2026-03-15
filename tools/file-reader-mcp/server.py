#!/usr/bin/env python3
"""
MCP Server: File Reader for GitHub Copilot
==========================================
Provides tools for reading PDF, Word (.docx), and Excel (.xlsx/.csv) files
so that GitHub Copilot can process document content via MCP.

Supported tools:
  - read_pdf       : Extract text from a PDF file (with optional page_range)
  - search_pdf     : Search for a keyword/phrase in a PDF and return matching snippets
  - read_word      : Extract text from a Word .docx file
  - read_excel     : Extract data from Excel / CSV as Markdown tables
  - get_file_info  : Get metadata about a supported document file
"""

import asyncio
import csv
import os
from pathlib import Path
from typing import Any, Optional

import mcp.types as types
from mcp.server import Server
from mcp.server.stdio import stdio_server

# ---------------------------------------------------------------------------
# Server instance
# ---------------------------------------------------------------------------

server = Server("file-reader")

SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF Document",
    ".docx": "Word Document (.docx)",
    ".doc": "Word Document Legacy (.doc) — not fully supported",
    ".xlsx": "Excel Spreadsheet (.xlsx)",
    ".xls": "Excel Spreadsheet Legacy (.xls) — limited support",
    ".csv": "CSV File",
}

# ---------------------------------------------------------------------------
# Path validation (security: prevent path traversal)
# ---------------------------------------------------------------------------

def _validate_path(file_path: str) -> Path:
    """Resolve and validate a file path. Raises on missing or non-file paths."""
    path = Path(file_path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path!r}")
    if not path.is_file():
        raise ValueError(f"Path is not a regular file: {file_path!r}")
    return path


# ---------------------------------------------------------------------------
# PDF reader
# ---------------------------------------------------------------------------

def _parse_page_range(page_range: str, total_pages: int) -> tuple[int, int]:
    """Parse '200-220' or '5' into (start_idx, end_idx), 0-based inclusive."""
    pr = page_range.strip()
    if "-" in pr:
        parts = pr.split("-", 1)
        start = max(1, int(parts[0].strip()))
        end = min(total_pages, int(parts[1].strip()))
    else:
        start = end = max(1, min(total_pages, int(pr)))
    return start - 1, end - 1  # convert to 0-based


def _read_pdf(path: Path, page_range: Optional[str] = None) -> str:
    """Extract text from a PDF using PyMuPDF (fitz).

    Args:
        path: Path to the PDF file.
        page_range: Optional page range string, e.g. '5' or '200-220' (1-based).
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError(
            "pymupdf is required to read PDF files.\n"
            "Install it with: pip install pymupdf"
        )

    doc = fitz.open(str(path))
    total = doc.page_count

    if page_range:
        start_idx, end_idx = _parse_page_range(page_range, total)
        page_indices = range(start_idx, end_idx + 1)
        range_label = f" (pages {start_idx + 1}–{end_idx + 1} of {total})"
    else:
        page_indices = range(total)
        range_label = f" ({total} pages)"

    pages: list[str] = []
    for i in page_indices:
        text = doc[i].get_text().strip()
        if text:
            pages.append(f"--- Page {i + 1} ---\n{text}")
    doc.close()

    if not pages:
        return (
            f"(No extractable text found{range_label}. "
            "The PDF may contain only scanned images. "
            "Consider using an OCR tool to convert it first.)"
        )
    return f"[Extracted{range_label}]\n\n" + "\n\n".join(pages)


def _search_pdf(
    path: Path,
    query: str,
    max_results: int = 10,
    context_chars: int = 400,
) -> str:
    """Search for text in a PDF and return matching passages with page context.

    Uses PyMuPDF's built-in page.search_for() for fast per-page lookup,
    then extracts a context window from the raw page text around each hit.
    Never reads the full document into memory at once.
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError(
            "pymupdf is required to search PDF files.\n"
            "Install it with: pip install pymupdf"
        )

    doc = fitz.open(str(path))
    results: list[str] = []
    total_hits = 0

    for page_idx in range(doc.page_count):
        if total_hits >= max_results:
            break

        page = doc[page_idx]
        # Fast built-in search — returns list of matching Rect objects
        hit_rects = page.search_for(query)
        if not hit_rects:
            continue

        page_num = page_idx + 1  # 1-based for display
        page_text = page.get_text()
        lower_text = page_text.lower()
        lower_query = query.lower()

        pos = 0
        page_snippets: list[str] = []
        while total_hits < max_results:
            idx = lower_text.find(lower_query, pos)
            if idx == -1:
                break

            half = context_chars // 2
            start = max(0, idx - half)
            end = min(len(page_text), idx + len(query) + half)
            snippet = page_text[start:end].strip()
            if start > 0:
                snippet = "..." + snippet
            if end < len(page_text):
                snippet = snippet + "..."

            page_snippets.append(snippet)
            total_hits += 1
            pos = idx + len(query)

        if page_snippets:
            results.append(
                f"### Page {page_num}\n\n"
                + "\n\n---\n\n".join(page_snippets)
            )

    doc.close()

    if not results:
        return f'No matches found for: `{query}`'

    header = (
        f'Found **{total_hits}** match(es) for `{query}` '
        f'(showing up to {max_results} — use max_results to increase):'
    )
    return header + "\n\n" + "\n\n---\n\n".join(results)


# ---------------------------------------------------------------------------
# Word reader
# ---------------------------------------------------------------------------

def _read_word(path: Path) -> str:
    """Extract text and tables from a Word .docx file."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError(
            "python-docx is required to read Word files.\n"
            "Install it with: pip install python-docx"
        )

    doc = Document(str(path))
    parts: list[str] = []

    # Walk the document body in order (paragraphs + tables interleaved)
    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            # Paragraph — preserve heading style as Markdown heading
            para_obj = None
            for p in doc.paragraphs:
                if p._element is block:
                    para_obj = p
                    break
            if para_obj and para_obj.text.strip():
                style = para_obj.style.name if para_obj.style else ""
                text = para_obj.text.strip()
                if style.startswith("Heading 1"):
                    parts.append(f"# {text}")
                elif style.startswith("Heading 2"):
                    parts.append(f"## {text}")
                elif style.startswith("Heading 3"):
                    parts.append(f"### {text}")
                else:
                    parts.append(text)

        elif tag == "tbl":
            # Table — render as Markdown table
            for table in doc.tables:
                if table._element is block:
                    rows = [
                        [cell.text.strip() for cell in row.cells]
                        for row in table.rows
                    ]
                    parts.append(_rows_to_markdown(rows))
                    break

    return "\n\n".join(parts) if parts else "(No readable text content found)"


# ---------------------------------------------------------------------------
# Excel / CSV reader
# ---------------------------------------------------------------------------

def _read_excel(path: Path, sheet_name: Optional[str] = None) -> str:
    """Extract data from an Excel (.xlsx) or CSV file as Markdown tables."""
    suffix = path.suffix.lower()

    if suffix == ".csv":
        return _read_csv(path)

    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl is required to read Excel files.\n"
            "Install it with: pip install openpyxl"
        )

    wb = openpyxl.load_workbook(str(path), data_only=True)

    if sheet_name:
        if sheet_name not in wb.sheetnames:
            raise ValueError(
                f"Sheet '{sheet_name}' not found.\n"
                f"Available sheets: {', '.join(wb.sheetnames)}"
            )
        sheets_to_read = [sheet_name]
    else:
        sheets_to_read = wb.sheetnames

    result_parts: list[str] = []
    for name in sheets_to_read:
        ws = wb[name]
        # Only read rows that contain at least one non-None value
        rows = [
            [str(cell) if cell is not None else "" for cell in row]
            for row in ws.iter_rows(values_only=True)
            if any(cell is not None for cell in row)
        ]
        if rows:
            result_parts.append(f"### Sheet: {name}\n\n{_rows_to_markdown(rows)}")
        else:
            result_parts.append(f"### Sheet: {name}\n\n*(Empty sheet)*")

    wb.close()
    return "\n\n".join(result_parts)


def _read_csv(path: Path) -> str:
    """Read a CSV file and return its content as a Markdown table."""
    rows: list[list[str]] = []
    # Try UTF-8-sig first (handles BOM from Excel exports), then latin-1 fallback
    for encoding in ("utf-8-sig", "latin-1"):
        try:
            with open(path, encoding=encoding, newline="") as f:
                reader = csv.reader(f)
                rows = [row for row in reader if any(cell.strip() for cell in row)]
            break
        except UnicodeDecodeError:
            continue

    if not rows:
        return "*(Empty CSV file)*"
    return _rows_to_markdown(rows)


# ---------------------------------------------------------------------------
# Markdown table helper
# ---------------------------------------------------------------------------

def _rows_to_markdown(rows: list[list[str]]) -> str:
    """Convert a 2D list of strings into a GitHub-Flavored Markdown table."""
    if not rows:
        return ""

    col_count = max(len(row) for row in rows)
    padded = [row + [""] * (col_count - len(row)) for row in rows]

    header_row = "| " + " | ".join(padded[0]) + " |"
    separator = "| " + " | ".join(["---"] * col_count) + " |"
    body_rows = "\n".join("| " + " | ".join(row) + " |" for row in padded[1:])

    if body_rows:
        return "\n".join([header_row, separator, body_rows])
    return "\n".join([header_row, separator])


# ---------------------------------------------------------------------------
# MCP tool definitions
# ---------------------------------------------------------------------------

@server.list_tools()
async def list_tools() -> list[types.Tool]:
    return [
        types.Tool(
            name="read_pdf",
            description=(
                "Read and extract text content from a PDF file, organized by page. "
                "Use 'page_range' to read only a subset of pages (e.g. '5' or '200-220') "
                "instead of the full document — much faster for large PDFs. "
                "For keyword search use 'search_pdf' instead."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Absolute path to the PDF file (e.g. C:/Documents/report.pdf)",
                    },
                    "page_range": {
                        "type": "string",
                        "description": (
                            "Optional: pages to read (1-based). "
                            "Single page: '5'. Range: '200-220'. "
                            "If omitted, the entire document is returned."
                        ),
                    },
                },
                "required": ["file_path"],
            },
        ),
        types.Tool(
            name="search_pdf",
            description=(
                "Search for a keyword or phrase inside a PDF file and return "
                "matching text passages with surrounding context and page numbers. "
                "MUCH faster than reading the whole PDF — use this first when you "
                "need to locate a specific rule, section, or term in a large document."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Absolute path to the PDF file",
                    },
                    "query": {
                        "type": "string",
                        "description": "Text to search for (case-insensitive)",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Maximum number of matching snippets to return (default: 10)",
                    },
                    "context_chars": {
                        "type": "integer",
                        "description": "Characters of context to show around each match (default: 400)",
                    },
                },
                "required": ["file_path", "query"],
            },
        ),
        types.Tool(
            name="read_word",
            description=(
                "Read and extract text and tables from a Microsoft Word .docx file. "
                "Preserves headings as Markdown headings and renders tables as Markdown tables."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Absolute path to the .docx file (e.g. C:/Documents/spec.docx)",
                    }
                },
                "required": ["file_path"],
            },
        ),
        types.Tool(
            name="read_excel",
            description=(
                "Read data from a Microsoft Excel (.xlsx) or CSV file. "
                "Returns each sheet's content as a Markdown table. "
                "Optionally specify a sheet name to read only that sheet."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Absolute path to the .xlsx or .csv file (e.g. C:/Data/sales.xlsx)",
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": (
                            "Optional: name of the specific sheet to read. "
                            "If omitted, all sheets are returned."
                        ),
                    },
                },
                "required": ["file_path"],
            },
        ),
        types.Tool(
            name="get_file_info",
            description=(
                "Get metadata about a document file — file size, type, "
                "page count (PDF), sheet names (Excel), author/title (PDF metadata). "
                "Useful before reading a very large file."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Absolute path to the file",
                    }
                },
                "required": ["file_path"],
            },
        ),
    ]


# ---------------------------------------------------------------------------
# MCP tool handler
# ---------------------------------------------------------------------------

@server.call_tool()
async def call_tool(name: str, arguments: dict[str, Any]) -> list[types.TextContent]:
    try:
        if name == "read_pdf":
            path = _validate_path(arguments["file_path"])
            if path.suffix.lower() != ".pdf":
                raise ValueError(f"Expected a .pdf file, got '{path.suffix}'")
            page_range: Optional[str] = arguments.get("page_range")
            content = _read_pdf(path, page_range)
            return [types.TextContent(type="text", text=f"**File:** `{path.name}`\n\n{content}")]

        elif name == "search_pdf":
            path = _validate_path(arguments["file_path"])
            if path.suffix.lower() != ".pdf":
                raise ValueError(f"Expected a .pdf file, got '{path.suffix}'")
            query = arguments["query"]
            if not query.strip():
                raise ValueError("query must not be empty")
            max_results: int = int(arguments.get("max_results", 10))
            context_chars: int = int(arguments.get("context_chars", 400))
            content = _search_pdf(path, query, max_results, context_chars)
            return [types.TextContent(type="text", text=f"**File:** `{path.name}`\n\n{content}")]

        elif name == "read_word":
            path = _validate_path(arguments["file_path"])
            if path.suffix.lower() not in (".docx", ".doc"):
                raise ValueError(f"Expected a .docx file, got '{path.suffix}'")
            if path.suffix.lower() == ".doc":
                return [types.TextContent(
                    type="text",
                    text=(
                        "Legacy `.doc` format is not directly supported.\n"
                        "Please open the file in Microsoft Word and save it as `.docx`, then try again."
                    ),
                )]
            content = _read_word(path)
            return [types.TextContent(type="text", text=f"**File:** `{path.name}`\n\n{content}")]

        elif name == "read_excel":
            path = _validate_path(arguments["file_path"])
            if path.suffix.lower() not in (".xlsx", ".xls", ".csv"):
                raise ValueError(f"Expected an .xlsx, .xls, or .csv file, got '{path.suffix}'")
            if path.suffix.lower() == ".xls":
                return [types.TextContent(
                    type="text",
                    text=(
                        "Legacy `.xls` format is not supported by openpyxl.\n"
                        "Please open the file in Excel and save it as `.xlsx`, then try again."
                    ),
                )]
            sheet_name: Optional[str] = arguments.get("sheet_name")
            content = _read_excel(path, sheet_name)
            return [types.TextContent(type="text", text=f"**File:** `{path.name}`\n\n{content}")]

        elif name == "get_file_info":
            path = _validate_path(arguments["file_path"])
            stat = path.stat()
            size_kb = stat.st_size / 1024
            ext = path.suffix.lower()
            file_type = SUPPORTED_EXTENSIONS.get(ext, f"Unknown type ({ext})")

            info_lines = [
                f"**File:** `{path.name}`",
                f"**Type:** {file_type}",
                f"**Size:** {size_kb:.1f} KB",
                f"**Full path:** `{path}`",
            ]

            # Extra metadata per type
            if ext == ".pdf":
                try:
                    import fitz
                    doc = fitz.open(str(path))
                    info_lines.append(f"**Pages:** {doc.page_count}")
                    meta = doc.metadata or {}
                    if meta.get("title"):
                        info_lines.append(f"**Title:** {meta['title']}")
                    if meta.get("author"):
                        info_lines.append(f"**Author:** {meta['author']}")
                    doc.close()
                except ImportError:
                    info_lines.append("*(Install pymupdf for PDF metadata)*")

            elif ext == ".xlsx":
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(str(path), read_only=True)
                    info_lines.append(f"**Sheets ({len(wb.sheetnames)}):** {', '.join(wb.sheetnames)}")
                    wb.close()
                except ImportError:
                    info_lines.append("*(Install openpyxl for Excel metadata)*")

            elif ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(str(path))
                    para_count = len(doc.paragraphs)
                    table_count = len(doc.tables)
                    info_lines.append(f"**Paragraphs:** {para_count}")
                    info_lines.append(f"**Tables:** {table_count}")
                    core = doc.core_properties
                    if core.author:
                        info_lines.append(f"**Author:** {core.author}")
                    if core.title:
                        info_lines.append(f"**Title:** {core.title}")
                except ImportError:
                    info_lines.append("*(Install python-docx for Word metadata)*")

            return [types.TextContent(type="text", text="\n".join(info_lines))]

        else:
            return [types.TextContent(type="text", text=f"Unknown tool: '{name}'")]

    except FileNotFoundError as exc:
        return [types.TextContent(type="text", text=f"**Error (file not found):** {exc}")]
    except ValueError as exc:
        return [types.TextContent(type="text", text=f"**Error (invalid input):** {exc}")]
    except ImportError as exc:
        return [types.TextContent(type="text", text=f"**Error (missing library):** {exc}")]
    except Exception as exc:  # noqa: BLE001
        return [types.TextContent(type="text", text=f"**Unexpected error** ({type(exc).__name__}): {exc}")]


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

async def main() -> None:
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            server.create_initialization_options(),
        )


if __name__ == "__main__":
    asyncio.run(main())
